from django.db.models.functions import TruncMonth
from django.db.models import Count
import datetime
from .ienum import FILETYPE
import subprocess
from .pdfreader import PdfReader
from django.contrib.auth.models import User, Group
from rest_framework import viewsets, views
from api.serializers import UserSerializer, GroupSerializer
from rest_framework.response import Response
from rest_framework import status
from rest_framework.exceptions import ParseError
# uload package
from rest_framework.parsers import MultiPartParser, FormParser
from .serializers import FileUploadSerializer, WordRecognitionSerializer, FileImageTerrorismUploadSerializer, FileVisionPornUploadSerializer, AudioFileUploadSerializer
from .models import FileUpload, WordRecognition, FileImageTerrorismUpload, FileVisionPornUpload
# Handle Image
from PIL import Image
from io import BytesIO
import json
from .video import video
from .ocr.chineseocr import OCR
from django.conf import settings
from .serializers import VideoFileUploadSerializer, OcrGeneralSerializer, OcrIDCardSerializer, AudioFileInspectionSerializer, ImageFileUploadSerializer, WordRecognitionInspectionSerializer, OcrDrivinglicenseSerializer, OcrVehiclelicenseSerializer, OcrBusinesslicenseSerializer, OcrBankcardSerializer, OcrHandWrittenSerializer, OcrVehicleplateSerializer, HistoryRecordSerializer, OcrBusinessCardSerializer
from .models import VideoFileUpload, AudioFileUpload, OcrGeneral, OcrIDCard, AudioFileInspection, ImageFileUpload, WordRecognitionInspection, OcrDrivinglicense, OcrVehiclelicense, OcrBusinesslicense, OcrBankcard, OcrHandWritten, OcrVehicleplate, HistoryRecord, OcrBusinessCard,HistoryHashRecord
import os
import shutil
import uuid
import cv2
from .kaldi.audios import audio
from .sensitives.sensitives import sensitiveClass
import wave
import contextlib
import codecs
import chardet
from django.core.files import File
from urllib.request import urlopen
from tempfile import NamedTemporaryFile
from pydub import AudioSegment
import filetype
import docx
from .filetype import FileType
import platform
import time
from .tasks import task_check_video_default, task_check_video_android, task_check_video_ios
from .crons import run_django_rq_task
import numpy as np
if(platform.system() == "Windows"):
    import win32com.client as wc
    import pythoncom
import hashlib
from .util import changeImgAngle,cutImgByWH
        

#对文件进行hash
def get_file_md5(f):
    m = hashlib.md5()
    while True:
        #如果不用二进制打开文件，则需要先编码
        #data = f.read(1024).encode('utf-8')
        data = f.read(1024)  #将文件分块读取
        if not data:
            break
        m.update(data)
    return m.hexdigest()

def get_two_float(f_str, n):
    f_str = str(f_str)      # f_str = '{}'.format(f_str) 也可以转换为字符串
    a, b, c = f_str.partition('.')
    c = (c+"0"*n)[:n]       # 如论传入的函数有几位小数，在字符串后面都添加n为小数0
    return ".".join([a, c])


def UpdateHistoryRecord(serializer, filetype, result, maxtype, violence, porn):
    file_id = serializer.id
    file_type = filetype
    screenshot_url = ""
    duration = ""

    if filetype == FILETYPE.Image.value:
        if result.get('file_name'):
            file_name = result['file_name']
        else:
            file_name = serializer.image.name.split('/')[1]
        file_url = settings.FILE_URL + serializer.image.url
    elif file_type == FILETYPE.Video.value:
        if result.get('file_name'):
            file_name = result['file_name']
        else:
            file_name = serializer.video.name.split('/')[1]
        file_url = settings.FILE_URL + serializer.video.url
        screenshot_url = result["screenshot_url"]
        duration = result["duration"]
    elif file_type == FILETYPE.Audio.value:
        if result.get('file_name'):
            file_name = result['file_name']
        else:
            file_name = serializer.speech.name.split('/')[1]
        file_url = settings.FILE_URL + serializer.speech.url
    elif file_type == FILETYPE.Text.value:
        if result.get('file_name'):
            file_name = result['file_name']
        else:
            file_name = serializer.text.name.split('/')[1]
        file_url = settings.FILE_URL + serializer.text.url
    elif file_type == FILETYPE.Content.value:
        file_name = ""
        file_url = ""
    else:
        file_name = "other"
        file_url = "other"

    inspection_result = result

    violence_percent = "0"
    violence_sensitivity_level = "-1"
    if violence is not None:
        violence_percent = get_two_float(float(violence), 2)
        if (float(violence) < settings.VIOLENCESCORE_MIN):
            violence_sensitivity_level = "0"
        if (float(violence) >= settings.VIOLENCESCORE_MIN and float(violence) <= settings.VIOLENCESCORE_MAX):
            violence_sensitivity_level = "1"
        if (float(violence) > settings.VIOLENCESCORE_MAX):
            violence_sensitivity_level = "2"

    porn_percent = "0"
    porn_sensitivity_level = "-1"
    if porn is not None:
        porn_percent = get_two_float(float(porn), 2)
        if (float(porn) < settings.PORNSCORE_MIN):
            porn_sensitivity_level = "0"
        if (float(porn) >= settings.PORNSCORE_MIN and float(porn) <= settings.PORNSCORE_MAX):
            porn_sensitivity_level = "1"
        if (float(porn) > settings.PORNSCORE_MAX):
            porn_sensitivity_level = "2"

    max_sensitivity_type = maxtype
    max_sensitivity_percent = "0.00"

    content = ""
    web_text = ""
    app_text = ""
    if maxtype == 'violence':
        max_sensitivity_level = violence_sensitivity_level
        max_sensitivity_percent = violence_percent
        if result.get('text_content'):
            content = result["text_content"]
        if result.get('sensitive_info'):
            web_text = result["sensitive_info"]["web_text"]
        if result.get('sensitive_info'):
            app_text = result["sensitive_info"]["app_text"]
    elif maxtype == 'porn':
        max_sensitivity_level = porn_sensitivity_level
        max_sensitivity_percent = porn_percent
        if result.get('text_content'):
            content = result["text_content"]
        if result.get('sensitive_info'):
            web_text = result["sensitive_info"]["web_text"]
        if result.get('sensitive_info'):
            app_text = result["sensitive_info"]["app_text"]
    elif maxtype == 'violence_porn':
        max_sensitivity_level = violence_sensitivity_level
        max_sensitivity_percent = violence_percent
        if result.get('text_content'):
            content = result["text_content"]
        if result.get('sensitive_info'):
            web_text = result["sensitive_info"]["web_text"]
        if result.get('sensitive_info'):
            app_text = result["sensitive_info"]["app_text"]
    elif maxtype == 'text' and file_type == FILETYPE.Text.value:
        max_sensitivity_level = None
        max_sensitivity_percent = "0.00"
        if result.get('text_content'):
            content = result["text_content"]
        if result.get('sensitive_info'):
            web_text = result["sensitive_info"]["web_text"]
        if result.get('sensitive_info'):
            app_text = result["sensitive_info"]["app_text"]
    elif maxtype == 'text' and file_type == FILETYPE.Content.value:
        max_sensitivity_level = None
        max_sensitivity_percent = "0.00"
        content = serializer.text
        web_text = result["web_text"]
        app_text = result["app_text"]
    elif maxtype == 'ocr':
        max_sensitivity_level = None
        max_sensitivity_percent = "0.00"
        content = result['content']
        web_text = result['text']
        inspection_result = result['content']
    elif maxtype == 'audio':
        max_sensitivity_level = None
        max_sensitivity_percent = "0.00"
        content = result['text']
    else:
        max_sensitivity_level = None
        max_sensitivity_percent = "0.00"

    process_status = 2
    system_id = serializer.system_id
    channel_id = serializer.channel_id
    user_id = serializer.user_id

    serial_number = int(time.time())
    if file_type == FILETYPE.Video.value and result.get('serial_number') is not None:
        serial_number = result["serial_number"]

    if file_type == FILETYPE.Video.value and result.get('status') is not None and result.get('status') == 3:
        process_status = 3

    draw_url = ""
    if result.get('draw_url') is not None:
        draw_url = result["draw_url"]

    HistoryRecord.objects.create(
        file_id=file_id, file_name=file_name,
        file_url=file_url, file_type=file_type,
        inspection_result=inspection_result, max_sensitivity_type=max_sensitivity_type,
        max_sensitivity_level=max_sensitivity_level, max_sensitivity_percent=max_sensitivity_percent, violence_percent=violence_percent,
        violence_sensitivity_level=violence_sensitivity_level, porn_percent=porn_percent,
        porn_sensitivity_level=porn_sensitivity_level, content=content,
        web_text=web_text, app_text=app_text, process_status=process_status,
        system_id=system_id, channel_id=channel_id, user_id=user_id,
        screenshot_url=screenshot_url, duration=duration, serial_number=serial_number,
        draw_url=draw_url
    )

def UpdateHistoryHashRecord(file_id, file_name, file_url, file_type, result,hash_value):
    file_id = file_id
    file_type = file_type
    file_name = file_name
    file_url = file_url
    inspection_result = result
    #保存校验后的hash值记录
    HistoryHashRecord.objects.create(
        file_id=file_id, file_name=file_name,
        file_url=file_url, file_type=file_type,
        inspection_result=inspection_result,
        hash_value=hash_value
    )


def RunShellWithReturnCode(command):
    p = subprocess.Popen(command, stdout=subprocess.PIPE,
                         stderr=subprocess.PIPE, shell=True)
    p.wait()
    output = ""
    error = ""
    while True:
        line = p.stdout.read()
        if not line:
            break
        output += line.decode("utf-8")

    while True:
        err = p.stderr.read()
        if not err:
            break
        error += err.decode("utf-8")
    return output, error


# 图片旋转
def rotate_bound(image, angle):
    # 获取宽高
    (h, w) = image.shape[:2]
    (cX, cY) = (w // 2, h // 2)

    # 提取旋转矩阵 sin cos
    M = cv2.getRotationMatrix2D((cX, cY), -angle, 1.0)
    cos = np.abs(M[0, 0])
    sin = np.abs(M[0, 1])

    # 计算图像的新边界尺寸
    nW = int((h * sin) + (w * cos))
    # nH = int((h * cos) + (w * sin))
    nH = h

    # 调整旋转矩阵
    M[0, 2] += (nW / 2) - cX
    M[1, 2] += (nH / 2) - cY

    return cv2.warpAffine(image, M, (nW, nH), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

class UserViewSet(viewsets.ModelViewSet):
    """
    API endpoint that allows users to be viewed or edited.
    """
    queryset = User.objects.all().order_by('-date_joined')
    serializer_class = UserSerializer


class GroupViewSet(viewsets.ModelViewSet):
    """
    API endpoint that allows groups to be viewed or edited.
    """
    queryset = Group.objects.all()
    serializer_class = GroupSerializer


class FileUploadViewSet(viewsets.ModelViewSet):

    queryset = FileUpload.objects.all()
    serializer_class = FileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        # file_obj = self.request.data.get('datafile')
        # print (file_obj)
        # ...
        # do some stuff with uploaded file
        # ...
        # try:
        #     img = Image.open(file_obj)
        #     # img.verify()
        #     pic_io = BytesIO()
        #     img.save(pic_io,img.format)

        # except:
        #     raise ParseError("Unsupported image type")

        file_path = iserializer.datafile.path
        check_result = settings.VIOLENCE.check_violence(file_path)
        # print (check_result)

        result = {
            "ret": 0,
            "msg": "ok",
            "data": {
                "tag_list": [
                    {
                        "tag_name": "protest",
                        "probability": check_result['protest']
                    },
                    {
                        "tag_name": "violence",
                        "probability": check_result['violence']
                    },
                    {
                        "tag_name": "sign",
                        "probability": check_result['sign']
                    },
                    {
                        "tag_name": "photo",
                        "probability": check_result['photo']
                    },
                    {
                        "tag_name": "fire",
                        "probability": check_result['fire']
                    },
                    {
                        "tag_name": "police",
                        "probability": check_result['police']
                    },
                    {
                        "tag_name": "children",
                        "probability": check_result['children']
                    },
                    {
                        "tag_name": "group_20",
                        "probability": check_result['group_20']
                    },
                    {
                        "tag_name": "group_100",
                        "probability": check_result['group_100']
                    },
                    {
                        "tag_name": "flag",
                        "probability": check_result['flag']
                    },
                    {
                        "tag_name": "night",
                        "probability": check_result['night']
                    },
                    {
                        "tag_name": "shouting",
                        "probability": check_result['shouting']
                    }]
            }
        }
        serializer.save(result=result)

        return Response(status=status.HTTP_201_CREATED)


class WordRecognitionViewSet(viewsets.ModelViewSet):

    queryset = WordRecognition.objects.all()
    serializer_class = WordRecognitionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()

        text = iserializer.text
        sensitive_list = sensitiveClass().check_sensitiveWords(text)

        if sensitive_list.get('sensitive_hit_flag') == 0:
            ret = 1
            msg = "无匹配记录"
            max_sensitivity_type = None
        else:
            ret = 0
            msg = "匹配到记录"
            max_sensitivity_type = 'text'

        data = sensitive_list
        serializer.save(ret=ret, msg=msg, data=data, text=iserializer.text)

        # 更新历史记录
        # UpdateHistoryRecord(iserializer, FILETYPE.Content.value,
        #                     data, max_sensitivity_type, None, None)

        return Response(status=status.HTTP_201_CREATED)


class WordRecognitionInspectionViewSet(viewsets.ModelViewSet):

    queryset = WordRecognitionInspection.objects.all()
    serializer_class = WordRecognitionInspectionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()

        # 增加网络URL文件上传
        if iserializer.text_url and not iserializer.text:
            txt_temp = NamedTemporaryFile(delete=True)
            txt_temp.write(urlopen(iserializer.text_url).read())
            txt_temp.flush()
            iserializer.text.save(os.path.basename(
                iserializer.text_url), File(txt_temp))

        # word格式文件读取
        filetype = FileType().filescanner(iserializer.text.path)
        if filetype is None:
            print('Cannot guess file type!')
            return Response(status=status.HTTP_400_BAD_REQUEST)

        text_content = ""
        sensitive_map = {}
        if filetype == 'zip':
            doc = docx.Document(iserializer.text.path)
            docText = '\n'.join(
                [paragraph.text for paragraph in doc.paragraphs])
            text_content = docText
        elif filetype == 'wps':
            # 仅支持windows
            if(platform.system() == "Windows"):
                pythoncom.CoInitialize()
                word = wc.Dispatch("Word.Application")
                doc = word.Documents.Open(iserializer.text.path)
                docx_path = iserializer.text.path.split(".doc")[0] + '.docx'
                doc.SaveAs(docx_path, 12)
                doc.Close
                word.Quit
                doc = docx.Document(docx_path)
                docText = '\n'.join(
                    [paragraph.text for paragraph in doc.paragraphs])
                text_content = docText
            else:
                # 仅支持ubuntu
                cmd = 'antiword -m UTF-8 ' + iserializer.text.path
                docText, errText = RunShellWithReturnCode(cmd)
                if len(errText) > 0:
                    text_content = 'doc文档内容过短，请重新上传'
                else:
                    text_content = docText
        elif filetype == 'pdf':
            pdfText = PdfReader().parse(iserializer.text.path)
            text_content = pdfText
        else:
            txtfile = iserializer.text.path
            # print(txtfile)
            # 增加gbk编码格式转换
            f_test = open(txtfile, 'rb')
            file_type = chardet.detect(f_test.read(100))

            if (file_type['encoding'] == 'GB2312'):
                f = codecs.open(txtfile, 'r', encoding='gbk', errors='ignore')
            elif (file_type['encoding'] == 'UTF-8-SIG'):
                f = codecs.open(txtfile, 'r', encoding='utf-8',
                                errors='ignore')
            elif (file_type['encoding'] == 'ascii'):
                f = codecs.open(txtfile, 'r', encoding='gbk', errors='ignore')
            else:
                f = codecs.open(txtfile, 'r', errors='ignore')

            try:
                for line in f:
                    text_content += line
            except Exception as e:
                print("The content get some error: " + line)
                print(e)
                msg = "内容获取异常"
                ret = 1
            else:
                print("Read content successfully!")

        sensitive_list = sensitiveClass().check_sensitiveWords(text_content)
        sensitive_map["text_content"] = text_content
        sensitive_map["sensitive_info"] = sensitive_list
        data = sensitive_map
        if sensitive_list.get('sensitive_hit_flag') == 0:
            ret = 1
            msg = "无匹配记录"
            max_sensitivity_type = None
        else:
            ret = 0
            msg = "匹配到记录"
            max_sensitivity_type = 'text'
        serializer.save(ret=ret, msg=msg, data=data, text=iserializer.text)

        # 更新历史记
        # UpdateHistoryRecord(iserializer, FILETYPE.Text.value,
        #                     data, max_sensitivity_type, None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrGeneralViewSet(viewsets.ModelViewSet):

    queryset = OcrGeneral.objects.all()
    serializer_class = OcrGeneralSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "通用OCR"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        dataArr = []
        boxArr = []
        for each in arr:
            dataArr.append(each["text"])
            boxArr.append(each["box"])
        result = {
            'content': dataArr,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        serializer.save(data=dataArr, ret=ret, msg=msg,box=boxArr,draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrIDCardViewSet(viewsets.ModelViewSet):

    queryset = OcrIDCard.objects.all()
    serializer_class = OcrIDCardSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "身份证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        for each in arr:
            name = ""
            if(each['name'] == '姓名'):
                name = "name"
                count = count + 1
            if(each['name'] == '性别'):
                name = "sex"
                count = count + 1
            if(each['name'] == '民族'):
                name = "nation"
                count = count + 1
            if(each['name'] == '出生年月'):
                name = "birth"
                count = count + 1
            if(each['name'] == '身份证号码'):
                name = "id"
                count = count + 1
            if(each['name'] == '身份证地址'):
                name = "address"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 3):
            ret = 1
            msg = "请上传身份证图片"
        serializer.save(data=dataMap, ret=ret, msg=msg, box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class FileImageTerrorismUploadViewSet(viewsets.ModelViewSet):
    queryset = FileImageTerrorismUpload.objects.all()
    serializer_class = FileImageTerrorismUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        check_result = settings.VIOLENCE.check_violence(file_path)

        #增加文本识别
        bill_model = "通用OCR"
        ocr_result = OCR().getWordRecognition(file_path, bill_model)
        arr = ocr_result['res']
        dataArr = []
        for each in arr:
            dataArr.append(each["text"])

        sensitive_list = sensitiveClass().check_sensitiveWords(ocr_result['text'])

        violence = check_result['violence']
        resultMap = {}
        resultMap['violence'] = get_two_float(float(violence) * 100, 2)
        resultMap['content'] = dataArr
        resultMap["text_content"] = ocr_result['text']
        resultMap['text'] = ocr_result['text']
        resultMap['sensitive_info'] = sensitive_list
        resultMap['web_text'] = sensitive_list['web_text']
        resultMap['file_name'] = self.request.FILES['image'].name

        serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)

        # 更新历史记录
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            resultMap, 'violence', resultMap['violence'], None)

        return Response(status=status.HTTP_201_CREATED)


class FileVisionPornUploadViewSet(viewsets.ModelViewSet):
    queryset = FileVisionPornUpload.objects.all()
    serializer_class = FileVisionPornUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        scores = settings.NSFW.caffe_preprocess_and_compute_api(file_path)

        #增加文本识别
        bill_model = "通用OCR"
        ocr_result = OCR().getWordRecognition(file_path, bill_model)
        arr = ocr_result['res']
        dataArr = []
        for each in arr:
            dataArr.append(each["text"])

        sensitive_list = sensitiveClass().check_sensitiveWords(ocr_result['text'])

        resultMap = {}
        resultMap['normal_hot_porn'] = get_two_float(float(scores[1]) * 100, 2)
        resultMap['content'] = dataArr
        resultMap["text_content"] = ocr_result['text']
        resultMap['text'] = ocr_result['text']
        resultMap['sensitive_info'] = sensitive_list
        resultMap['web_text'] = sensitive_list['web_text']
        resultMap['file_name'] = self.request.FILES['image'].name
        # print (check_result)
        serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)

        # 更新历史记录
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            resultMap, 'porn', None, resultMap['normal_hot_porn'])

        return Response(status=status.HTTP_201_CREATED)


class VideoFileUploadViewSet(viewsets.ModelViewSet):
    queryset = VideoFileUpload.objects.all()
    serializer_class = VideoFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def list(self, request):
        # 获取参数
        requestData = request.query_params
        object_id = requestData.get('id')
        is_task = requestData.get('is_task')


        # 根据条件过滤
        conditions = {}
        if object_id is not None:
            conditions['id'] = object_id
        
        if is_task is not None and is_task == '1' and object_id:
            iVideoFileUpload = VideoFileUpload.objects.get(id=object_id)
            data = eval(iVideoFileUpload.data)
            file_path = iVideoFileUpload.video.path
            orientation = iVideoFileUpload.orientation
            serial_number = data['serial_number']
            # 增加识别中状态
            iHistoryRecord = HistoryRecord.objects.get(serial_number=serial_number)
            iHistoryRecord.process_status = 1
            iHistoryRecord.save()

            #对视频文件进行hash，判断文件是否已经上传过
            file_md5 = ''
            with open(file_path, 'rb') as f:
                file_md5 = get_file_md5(f)
                historyHashRecord = HistoryHashRecord.objects.filter(hash_value=file_md5)
                if historyHashRecord.exists():
                    resultMap = json.loads(historyHashRecord[0].inspection_result.replace("'","\""))
                else:
                    #resultMap = video().check_video_V2(file_path, orientation, serial_number)
                    resultMap = video().check_video_imgs_similarity_filter(file_path, orientation, serial_number)
                    #保存hash值记录
                    file_id = iHistoryRecord.file_id
                    file_type = FILETYPE.Video.value
                    file_name = iHistoryRecord.file_name
                    file_url = iHistoryRecord.file_url
                    UpdateHistoryHashRecord(file_id, file_name, file_url, file_type, resultMap,file_md5)
            
            iVideoFileUpload.data = resultMap
            iVideoFileUpload.serial_number = data['serial_number']
            iVideoFileUpload.video_url = data['video_url']
            iVideoFileUpload.is_task = 1
            iVideoFileUpload.sync = 1
            iVideoFileUpload.ret = 0
            iVideoFileUpload.msg = "成功"
            iVideoFileUpload.save()

            max_sensitivity_percent = "0.00"
            maxtype = resultMap['max_sensitivity_type']
            if maxtype == 'violence':
                max_sensitivity_percent = resultMap['violence_percent']
            elif maxtype == 'porn':
                max_sensitivity_percent = resultMap['porn_percent']
            elif maxtype == 'violence_porn':
                max_sensitivity_percent = resultMap['violence_percent']

            # iHistoryRecord = HistoryRecord.objects.get(serial_number=serial_number)
            iHistoryRecord.inspection_result = resultMap
            iHistoryRecord.max_sensitivity_type = resultMap['max_sensitivity_type']
            iHistoryRecord.max_sensitivity_level = resultMap['max_sensitivity_level']
            iHistoryRecord.max_sensitivity_percent = max_sensitivity_percent
            iHistoryRecord.violence_percent = resultMap['violence_percent']
            iHistoryRecord.violence_sensitivity_level = resultMap['violence_sensitivity_level']
            iHistoryRecord.porn_percent = resultMap['porn_percent']
            iHistoryRecord.porn_sensitivity_level = resultMap['porn_sensitivity_level']
            iHistoryRecord.content = ""
            iHistoryRecord.web_text = ""
            iHistoryRecord.app_text = ""
            iHistoryRecord.process_status = 2
            iHistoryRecord.screenshot_url = resultMap["screenshot_url"]
            iHistoryRecord.duration = resultMap["duration"]
            iHistoryRecord.save()

        queryset = VideoFileUpload.objects.filter(**conditions)

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    def perform_create(self, serializer):
        iserializer = serializer.save()

        # 增加网络URL文件上传
        if iserializer.video_url and not iserializer.video:
            video_temp = NamedTemporaryFile(delete=True)
            video_temp.write(urlopen(iserializer.video_url).read())
            video_temp.flush()
            iserializer.video.save(os.path.basename(
                iserializer.video_url), File(video_temp))

        file_path = iserializer.video.path
        orientation = iserializer.orientation
        sync = iserializer.sync
        system_id = iserializer.system_id
        serial_number = int(time.time())
        if iserializer.screenshot.name != None:
            screenshot_file_path = iserializer.screenshot.path
            screenshot_file = cv2.imread(screenshot_file_path)
            # 增加图片旋转矫正
            # if orientation:
            #     # Flipped Horizontally 水平翻转
            #     if orientation == 1:
            #         screenshot_file = rotate_bound(screenshot_file, 90.000)
            #     elif orientation == 3:
            #         screenshot_file = rotate_bound(screenshot_file, -90.000)
            #     elif orientation == 6:
            #         screenshot_file = rotate_bound(screenshot_file, 180.000)
            #     cv2.imwrite(screenshot_file_path, screenshot_file)
            
        else:
            screenshot_file_path = ""
        
        if sync or sync is None:
            #对视频文件进行hash，判断文件是否已经上传过
            file_md5 = ''
            with open(file_path, 'rb') as f:
                file_md5 = get_file_md5(f)
                historyHashRecord = HistoryHashRecord.objects.filter(hash_value=file_md5)
                if historyHashRecord.exists():
                    ret = 0
                    msg = "成功"
                    resultMap = json.loads(historyHashRecord[0].inspection_result.replace("'","\""))
                    resultMap['file_name'] = self.request.FILES['video'].name
                    serializer.save(data=resultMap, ret=ret,
                                    msg=msg, video=iserializer.video)
                else:
                    start = time.clock()
                    #resultMap = video().check_video_V2(file_path, orientation, serial_number)

                    resultMap = video().check_video_imgs_similarity_filter(file_path, orientation, serial_number)

                    elapsed = (time.clock() - start)
                    print("Time used:",elapsed)

                    ret = 0
                    msg = "成功"
                    serializer.save(data=resultMap, ret=ret,
                                    msg=msg, video=iserializer.video)
                    #保存hash值记录
                    file_id = iserializer.id
                    file_type = FILETYPE.Video.value
                    file_name = iserializer.video.name.split('/')[1]
                    file_url = settings.FILE_URL + iserializer.video.url
                    resultMap['file_name'] = self.request.FILES['video'].name
                    UpdateHistoryHashRecord(file_id, file_name, file_url, file_type, resultMap,file_md5)

                # 更新历史记录
                UpdateHistoryRecord(iserializer, FILETYPE.Video.value,
                    resultMap, resultMap['max_sensitivity_type'],
                    resultMap['violence_percent'], resultMap['porn_percent'])
        else:
            ret = 0
            msg = "成功"
            resultMap = {}
            p, f = os.path.split(file_path)
            if screenshot_file_path != "":
                sp,sf = os.path.split(screenshot_file_path)
                resultMap['screenshot_url'] = settings.VIDEO_URL + sf
            else:
                resultMap['screenshot_url'] = ""
            resultMap['video_url'] = settings.VIDEO_URL + f
            resultMap['violence_sensitivity_level'] = "-1"
            resultMap['porn_sensitivity_level'] = "-1"
            resultMap['video_evidence_information'] = []
            resultMap['violence_evidence_information'] = []
            resultMap['porn_evidence_information'] = []
            resultMap['interval'] = ""
            resultMap['duration'] = 0
            resultMap['fps'] = 0
            resultMap['taketimes'] = 0
            resultMap['max_sensitivity_type'] = "-1"
            resultMap['max_sensitivity_level'] = "-1"
            resultMap['max_sensitivity_percent'] = "0.00"
            
            resultMap['serial_number'] = serial_number
            resultMap['progress'] = "50%"
            resultMap['status'] = 3
            resultMap['file_name'] = self.request.FILES['video'].name

            serializer.save(data=resultMap, ret=ret,
                            msg=msg, video=iserializer.video, screenshot=iserializer.screenshot)

            # 更新历史记录
            UpdateHistoryRecord(iserializer, FILETYPE.Video.value,
                                resultMap, resultMap['max_sensitivity_type'],
                                None, None)
            
            # 上传成功，并创建识别任务
            if system_id == 2:
                task_check_video_android.delay(iserializer, serial_number)
            elif system_id == 3:
                task_check_video_ios.delay(iserializer, serial_number)
            else:
                task_check_video_default.delay(iserializer, serial_number)

            # 执行任务
            if settings.IS_SUPPORT_RQ:
                run_django_rq_task()

        return Response(status=status.HTTP_201_CREATED)


class AudioFileUploadViewSet(viewsets.ModelViewSet):
    queryset = AudioFileUpload.objects.all()
    serializer_class = AudioFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.speech_url and not iserializer.speech:
            speech_temp = NamedTemporaryFile(delete=True)
            speech_temp.write(urlopen(iserializer.speech_url).read())
            speech_temp.flush()
            iserializer.speech.save(os.path.basename(
                iserializer.speech_url), File(speech_temp))

        file_path = iserializer.speech.path
        size = os.path.getsize(file_path)
        if size <= 44:
            check_result = '录音时间太短，请重新录音！'
        else:
            check_result = audio().getOneAudioContent(file_path)
        # print (check_result)
        resultMap = {}
        resultMap['text'] = check_result
        serializer.save(data=resultMap, ret=ret, msg=msg,
                        speech=iserializer.speech)

        # 更新历史记录
        UpdateHistoryRecord(iserializer, FILETYPE.Audio.value,
                            resultMap, 'audio', None, None)

        return Response(status=status.HTTP_201_CREATED)


class AudioFileInspectionViewSet(viewsets.ModelViewSet):
    queryset = AudioFileInspection.objects.all()
    serializer_class = AudioFileInspectionSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.speech_url and not iserializer.speech:
            speech_temp = NamedTemporaryFile(delete=True)
            speech_temp.write(urlopen(iserializer.speech_url).read())
            speech_temp.flush()
            iserializer.speech.save(os.path.basename(
                iserializer.speech_url), File(speech_temp))

        # 转换mp3-wav
        kind = filetype.guess(iserializer.speech.path)
        if kind is None:
            print('Cannot guess file type!')
            return Response(status=status.HTTP_400_BAD_REQUEST)

        print(kind.extension)
        file_path = ''
        if kind.extension == 'mp3':
            sound = AudioSegment.from_mp3(iserializer.speech.path)
            destin_path = iserializer.speech.path.split(".mp3")[0] + '.wav'
            sound.export(destin_path, format='wav')
            file_path = destin_path
        elif kind.extension == 'ogg':
            sound = AudioSegment.from_ogg(iserializer.speech.path)
            destin_path = iserializer.speech.path.split(".ogg")[0] + '.wav'
            sound.export(destin_path, format='wav')
            file_path = destin_path
        else:
            file_path = iserializer.speech.path

        if len(file_path) > 0:
            duration = 0
            with contextlib.closing(wave.open(file_path, 'r')) as f:
                frames = f.getnframes()
                rate = f.getframerate()
                duration = frames / float(rate)
            audio_content = audio().getOneAudioContent(file_path)
            check_result = sensitiveClass().check_sensitiveWords(audio_content)
            resultMap = {}
            resultMap["speech_time"] = duration
            resultMap["speech_contents"] = check_result
            serializer.save(data=resultMap, ret=ret, msg=msg,
                            speech=iserializer.speech)

            # 更新历史记录
            resultMap['text'] = check_result
            resultMap['file_name'] = self.request.FILES['speech'].name
            UpdateHistoryRecord(iserializer, FILETYPE.Audio.value,
                                resultMap, 'audio', None, None)

            return Response(status=status.HTTP_201_CREATED)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


class ImageFileUploadViewSet(viewsets.ModelViewSet):
    queryset = ImageFileUpload.objects.all()
    serializer_class = ImageFileUploadSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):
        iserializer = serializer.save()
        ret = 0
        msg = "成功"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        scores = settings.NSFW.caffe_preprocess_and_compute_api(file_path)
        resultMap = {}
        porn_sensitivity_level = "-1"
        porn_pencent = float(scores[1]) * 100
        if (porn_pencent < settings.PORNSCORE_MIN):
            porn_sensitivity_level = "0"
        if (porn_pencent >= settings.PORNSCORE_MIN and porn_pencent <= settings.PORNSCORE_MAX):
            porn_sensitivity_level = "1"
        if (porn_pencent > settings.PORNSCORE_MAX):
            porn_sensitivity_level = "2"
        resultMap['porn_sensitivity_level'] = porn_sensitivity_level
        resultMap['porn_percent'] = get_two_float(float(scores[1]) * 100, 2)

        check_result = settings.VIOLENCE.check_violence(file_path)
        violence = check_result['violence']
        violence_sensitivity_level = "-1"
        violence_pencent = float(violence) * 100
        if (violence_pencent < settings.VIOLENCESCORE_MIN):
            violence_sensitivity_level = "0"
        if (violence_pencent >= settings.VIOLENCESCORE_MIN and violence_pencent <= settings.VIOLENCESCORE_MAX):
            violence_sensitivity_level = "1"
        if (violence_pencent > settings.VIOLENCESCORE_MAX):
            violence_sensitivity_level = "2"
        resultMap['violence_sensitivity_level'] = violence_sensitivity_level
        resultMap['violence_percent'] = get_two_float(float(violence) * 100, 2)

        resultMap['politics_sensitivity_level'] = ""
        resultMap['politics_percent'] = ""
        resultMap['public_character_level'] = ""
        resultMap['public_percent'] = ""

         #增加文本识别
        bill_model = "通用OCR"
        ocr_result = OCR().getWordRecognition(file_path, bill_model)
        arr = ocr_result['res']
        dataArr = []
        for each in arr:
            dataArr.append(each["text"])

        sensitive_list = sensitiveClass().check_sensitiveWords(ocr_result['text'])

        resultMap['content'] = dataArr
        resultMap["text_content"] = ocr_result['text']
        resultMap['text'] = ocr_result['text']
        resultMap['sensitive_info'] = sensitive_list
        resultMap['web_text'] = sensitive_list['web_text']
        if iserializer.image_url:
            resultMap['file_name'] = iserializer.image.name.lstrip('photos/')
            image_url = settings.FILE_URL +  settings.MEDIA_URL + 'photos' + '/' + resultMap['file_name']
            serializer.save(data=resultMap, ret=ret,image_url=image_url,
                        msg=msg, image=iserializer.image)
        else:
            resultMap['file_name'] = self.request.FILES['image'].name
            serializer.save(data=resultMap, ret=ret,
                        msg=msg, image=iserializer.image)

        # 更新历史记录
        if float(violence) > float(scores[1]):
            max_sensitivity_type = 'violence'
        elif float(violence) < float(scores[1]):
            max_sensitivity_type = 'porn'
        else:
            max_sensitivity_type = 'violence_porn'

        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            resultMap, max_sensitivity_type, resultMap['violence_percent'], resultMap['porn_percent'])

        return Response(status=status.HTTP_201_CREATED)


class OcrDrivinglicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrDrivinglicense.objects.all()
    serializer_class = OcrDrivinglicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "驾驶证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        dataMap["license_type"] = ""
        dataMap["card_id"] = ""
        dataMap["driver"] = ""
        dataMap["sex"] = ""
        dataMap["nationality"] = ""
        dataMap["address"] = ""
        dataMap["birthday"] = ""
        dataMap["issue_date"] = ""
        dataMap["be_class"] = ""
        dataMap["valid_start"] = ""
        dataMap["valid_end"] = ""
        dataMap["after_five"] = ""
        dataMap["remark"] = ""

        for each in arr:
            name = ""
            if(each['name'] == '类型'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '证号'):
                name = "card_id"
                count = count + 1
            if(each['name'] == '姓名'):
                name = "driver"
                count = count + 1
            if(each['name'] == '性别'):
                name = "sex"
                count = count + 1
            if(each['name'] == '国籍'):
                name = "nationality"
                count = count + 1
            if(each['name'] == '住址'):
                name = "address"
                count = count + 1
            if(each['name'] == '出生日期'):
                name = "birthday"
                count = count + 1
            if(each['name'] == '初次领证日期'):
                name = "issue_date"
                count = count + 1
            if(each['name'] == '准驾车型'):
                name = "be_class"
                count = count + 1
            if(each['name'] == '有效起始日期'):
                name = "valid_start"
                count = count + 1
            if(each['name'] == '有效截止日期'):
                name = "valid_end"
                count = count + 1

            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        # if (len(arr) == 0 or count < 1):
        if(len(dataMap) <= 4 or dataMap['license_type'] != '中华人民共和国机动车驾驶证'):
            ret = 1
            msg = "请上传驾驶证图片"

        serializer.save(data=dataMap, ret=ret, msg=msg,box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrVehiclelicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrVehiclelicense.objects.all()
    serializer_class = OcrVehiclelicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "行驶证"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        dataMap["license_type"] = ""
        dataMap["plate_no"] = ""
        dataMap["vehicle_type"] = ""
        dataMap["owner"] = ""
        dataMap["address"] = ""
        dataMap["use_character"] = ""
        dataMap["model"] = ""
        dataMap["vin"] = ""
        dataMap["engine_no"] = ""
        dataMap["register_date"] = ""
        dataMap["issue_date"] = ""

        for each in arr:
            name = ""
            if(each['name'] == '类型'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '号牌号码'):
                name = "plate_no"
                count = count + 1
            if(each['name'] == '车辆类型'):
                name = "vehicle_type"
                count = count + 1
            if(each['name'] == '所有人'):
                name = "owner"
                count = count + 1
            if(each['name'] == '住址'):
                name = "address"
                count = count + 1
            if(each['name'] == '使用性质'):
                name = "use_character"
                count = count + 1
            if(each['name'] == '品牌型号'):
                name = "model"
                count = count + 1
            if(each['name'] == '车辆识别代号'):
                name = "vin"
                count = count + 1
            if(each['name'] == '发动机号码'):
                name = "engine_no"
                count = count + 1
            if(each['name'] == '注册日期'):
                name = "register_date"
                count = count + 1
            if(each['name'] == '发证日期'):
                name = "issue_date"
                count = count + 1
            dataMap[name] = each['text']
            #dataMap[each['name']] = each['text']
        #result = check_result
        # if (len(arr) == 0 or count < 1):
        if(len(dataMap) <= 4 or dataMap['license_type'] != '中华人民共和国机动车行驶证'):
            ret = 1
            msg = "请上传行驶证图片"

        serializer.save(data=dataMap, ret=ret, msg=msg, box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrBusinesslicenseViewSet(viewsets.ModelViewSet):

    queryset = OcrBusinesslicense.objects.all()
    serializer_class = OcrBusinesslicenseSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "营业执照"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        dataMap["license_type"] = ""
        dataMap["business_id"] = ""
        dataMap["business_name"] = ""
        dataMap["business_type"] = ""
        dataMap["address"] = ""
        dataMap["operator"] = ""
        dataMap["registered_capital"] = ""
        dataMap["register_date"] = ""
        dataMap["business_term"] = ""
        dataMap["scope"] = ""

        for each in arr:
            name = ""
            if(each['name'] == '营业执照'):
                name = "license_type"
                count = count + 1
            if(each['name'] == '统一社会信用代码'):
                name = "business_id"
                count = count + 1
            if(each['name'] == '名称'):
                name = "business_name"
                count = count + 1
            if(each['name'] == '类型'):
                name = "business_type"
                count = count + 1
            if(each['name'] == '住所'):
                name = "address"
                count = count + 1
            if(each['name'] == '法定代表人'):
                name = "operator"
                count = count + 1
            if(each['name'] == '注册资本'):
                name = "registered_capital"
                count = count + 1
            if(each['name'] == '成立日期'):
                name = "register_date"
                count = count + 1
            if(each['name'] == '营业期限'):
                name = "business_term"
                count = count + 1
            if(each['name'] == '经营范围'):
                name = "scope"
            dataMap[name] = each['text']
        if dataMap["license_type"] != "营业执照" and len(dataMap) >= 8:
            dataMap["license_type"] = "营业执照"

        if len(dataMap) <= 4:
            ret = 1
            msg = "请上传营业执照图片"
        serializer.save(data=dataMap, ret=ret, msg=msg, box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrBankcardViewSet(viewsets.ModelViewSet):

    queryset = OcrBankcard.objects.all()
    serializer_class = OcrBankcardSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "银行卡"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        dataMap["bank_name"] = ""
        dataMap["bank_cardno"] = ""
        dataMap["expiry_date"] = ""
        dataMap["card_type"] = ""
        dataMap["card_name"] = ""

        for each in arr:
            name = ""
            if(each['name'] == '银行名称'):
                name = "bank_name"
                count = count + 1
            if(each['name'] == '卡号'):
                name = "bank_cardno"
                count = count + 1
            if(each['name'] == '卡类型'):
                name = "card_type"
                count = count + 1
            if(each['name'] == '有效期'):
                name = "expiry_date"
                count = count + 1
            if(each['name'] == '卡名称'):
                name = "card_name"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 1):
            ret = 1
            msg = "请上传银行卡图片"

        serializer.save(data=dataMap, ret=ret, msg=msg, box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text']
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrHandWrittenViewSet(viewsets.ModelViewSet):

    queryset = OcrHandWritten.objects.all()
    serializer_class = OcrHandWrittenSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "手写体"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        #霍夫矫正旋转图片角度
        img = cv2.imread(file_path)
        oH ,oW= img.shape[0:2]

        # print (file_path)
        # check_result = OCR().getWordRecognition(file_path, bill_model)
        from handwrite.handwrite import HandWrite
        check_result = HandWrite().getWord(file_path)
        #识别后还原图片角度
        angle = check_result['angle']
        if angle != 0:
            splitStr = file_path.split(".")
            drawedImgPath = splitStr[0]+"_drawed."+splitStr[1]
            dnW,dnH = changeImgAngle(drawedImgPath,-angle)
            nW,nH = changeImgAngle(file_path,-angle)
            cutImgByWH(oW,oH,dnW,dnH,drawedImgPath)
            cutImgByWH(oW,oH,nW,nH,file_path)
     
        # print (check_result)
        arr = check_result['data']
        drawUrl = check_result['drawUrl']
        dataArr = []
        dataBox = []
        dataMap = {}
        dataMap["handwritten_content"] = ""
        for each in arr:
            if len(arr) >= 0:
                dataArr.append(each["text"])
                dataBox.append(each["box"])

        dataMap["handwritten_content"] = dataArr

        # result = check_result
        serializer.save(data=dataMap, ret=ret, msg=msg, box=dataBox, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': dataBox,
            'draw_url': drawUrl,
            'text': check_result['data'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrVehicleplateViewSet(viewsets.ModelViewSet):

    queryset = OcrVehicleplate.objects.all()
    serializer_class = OcrVehicleplateSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "车牌"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        # arr = check_result['res']
        # carPlateIdentity = CarPlateIdentity()
        # ret, car_num = carPlateIdentity.car_plate_identity(file_path)
        ret, car_num,box,drawUrl = settings.CARPLATEIDENTITY.car_plate_identity(file_path)
        dataMap = {}
        check_result = {}
        # count = 0
        if ret == False:
            ret = 1
            msg = "请上传车牌图片"
            dataMap['plate_no'] = "请上传车牌图片"
            check_result['text'] = "请上传车牌图片"
        else:
            dataMap['plate_no'] = car_num
            check_result['text'] = car_num

        # dataMap["plate_no"] = ""

        # for each in arr:
        #     name = ""
        #     if(each['name'] == '车牌号'):
        #         name = "plate_no"
        #         count = count + 1
        #     dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        # if (len(arr) == 0 or count < 1 or dataMap["plate_no"] == "其他"):
        #     ret = 1
        #     msg = "请上传车牌图片"
        serializer.save(data=dataMap, ret=ret, msg=msg, box=box, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': box,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class OcrBusinessCardViewSet(viewsets.ModelViewSet):

    queryset = OcrBusinessCard.objects.all()
    serializer_class = OcrBusinessCardSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def perform_create(self, serializer):

        iserializer = serializer.save()
        ret = 0
        msg = "成功"
        bill_model = "名片"

        # 增加网络URL文件上传
        if iserializer.image_url and not iserializer.image:
            img_temp = NamedTemporaryFile(delete=True)
            img_temp.write(urlopen(iserializer.image_url).read())
            img_temp.flush()
            iserializer.image.save(os.path.basename(
                iserializer.image_url), File(img_temp))

        file_path = iserializer.image.path
        # print (file_path)
        check_result = OCR().getWordRecognition(file_path, bill_model)
        arr = check_result['res']
        drawUrl = check_result['drawUrl']
        com_arr = check_result['com_res']
        dataMap = {}
        boxArr = []
        count = 0

        #增加box
        for each in com_arr:
            boxArr.append(each["box"])

        dataMap["business_name"] = ""
        dataMap["position"] = ""
        dataMap["company"] = ""
        dataMap["address"] = ""
        dataMap["email"] = ""
        dataMap["phone"] = ""
        dataMap["telephone"] = ""
        dataMap["qq"] = ""
        dataMap["webchat"] = ""

        for each in arr:
            name = ""
            if(each['name'] == '姓名'):
                name = "business_name"
                count = count + 1
            if(each['name'] == '职位'):
                name = "position"
                count = count + 1
            if(each['name'] == '公司'):
                name = "company"
                count = count + 1
            if(each['name'] == '地址'):
                name = "address"
                count = count + 1
            if(each['name'] == '邮箱'):
                name = "email"
                count = count + 1
            if(each['name'] == '手机'):
                name = "phone"
                count = count + 1
            if(each['name'] == '电话'):
                name = "telephone"
                count = count + 1
            if(each['name'] == 'QQ'):
                name = "qq"
                count = count + 1
            if(each['name'] == '微信'):
                name = "webchat"
                count = count + 1
            dataMap[name] = each['text']
            # dataMap[each['name']] = each['text']
        # result = check_result
        if (len(arr) == 0 or count < 1 or dataMap["business_name"] == "其他"):
            ret = 1
            msg = "请上传名片图片"
        serializer.save(data=dataMap, ret=ret, msg=msg, box=boxArr, draw_url=drawUrl,
                        image=iserializer.image)

        # 更新历史记
        result = {
            'content': dataMap,
            'box': boxArr,
            'draw_url': drawUrl,
            'text': check_result['text'],
            'file_name': self.request.FILES['image'].name
        }
        UpdateHistoryRecord(iserializer, FILETYPE.Image.value,
                            result, 'ocr', None, None)

        return Response(status=status.HTTP_201_CREATED)


class HistoryRecordViewSet(viewsets.ModelViewSet):

    queryset = HistoryRecord.objects.all()
    serializer_class = HistoryRecordSerializer
    parser_classes = (MultiPartParser, FormParser,)

    def destroy(self, request, *args, **kwargs):
        # 获取参数
        requestData = request.query_params
        objecId = requestData.get("id")
        ids = requestData.get('ids')
        user_id = requestData.get('user_id')
        group_type = requestData.get('group_type')

        # 根据条件过滤
        conditions = {}
        if objecId is not None:
            conditions['id'] = objecId

        if ids is not None:
            ids = ids.split(',')
            ids_int = list(map(int, ids))
            conditions['id__in'] = ids_int

        if user_id is not None:
            conditions['user_id'] = user_id

        if group_type is not None:
            channel_ids = []
            if int(group_type) == 0:
                channel_ids = [1, 2, 3, 5, 6, 7, 8, 9, 10, 11, 12, 13, 15]
            elif int(group_type) == 1:
                channel_ids = [2, 3]
            elif int(group_type) == 2:
                channel_ids = [5, 6, 7, 8, 9, 10, 11, 12, 13]
            elif int(group_type) == 3:
                channel_ids = [15]
            else:
                channel_ids = [4, 14, 99]
            conditions['channel_id__in'] = channel_ids

        instance = HistoryRecord.objects.filter(**conditions)

        # instance = self.get_object()
        self.perform_destroy(instance)
        return Response(status=status.HTTP_204_NO_CONTENT)

    def retrieve(self, request, pk=None):
        # 获取实例
        historyRecord = self.get_object()
        # 序列化
        serializer = self.get_serializer(historyRecord)
        dataMap = {}
        dataMap['ret'] = 0
        dataMap['msg'] = "成功"
        dataMap['data'] = serializer.data
        return Response(data=dataMap, status=status.HTTP_200_OK)

    def list(self, request):
        # 获取参数
        requestData = request.query_params
        objecId = requestData.get('id')
        system_id = requestData.get('system_id')
        channel_id = requestData.get('channel_id')
        user_id = requestData.get('user_id')
        begin_time = requestData.get('begin_time')
        end_time = requestData.get('end_time')
        file_name = requestData.get('file_name')
        file_type = requestData.get('file_type')
        is_group = requestData.get('is_group')
        query_date = requestData.get('query_date')
        group_type = requestData.get('group_type')
        serial_number = requestData.get('serial_number')


        # 根据条件过滤
        conditions = {}
        if objecId is not None:
            iHistoryRecord = HistoryRecord.objects.get(id=objecId)
            iHistoryRecord.inspection_result = eval(iHistoryRecord.inspection_result)
            serializer = self.get_serializer(iHistoryRecord, many=False)
            if system_id is not None and system_id=='2':
                dataMap = {}
                dataMap['ret'] = 0
                dataMap['msg'] = "成功"
                dataMap['results'] = serializer.data
                dataMap['results']['video_evidence_information'] = serializer.data['inspection_result']['video_evidence_information']
                dataMap['results']['inspection_result'] = {}
                return Response(dataMap)
            else:
                dataMap = {}
                dataMap['ret'] = 0
                dataMap['msg'] = "成功"
                dataMap['results'] = serializer.data
                return Response(dataMap)

        if system_id is not None:
            conditions['system_id'] = system_id

        if channel_id is not None:
            conditions['channel_id'] = channel_id

        if user_id is not None:
            conditions['user_id'] = user_id

        if file_name is not None:
            conditions['file_name__contains'] = file_name

        if file_type is not None:
            conditions['file_type'] = file_type

        if begin_time is not None:
            begin_time_date = datetime.datetime.strptime(
                begin_time, "%Y-%m-%d %H:%M:%S")
            conditions['upload_time__gte'] = begin_time_date

        if end_time is not None:
            end_time_date = datetime.datetime.strptime(
                end_time, "%Y-%m-%d %H:%M:%S")
            conditions['upload_time__lte'] = end_time_date

        if serial_number is not None:
            iHistoryRecord = HistoryRecord.objects.get(serial_number=serial_number)
            iHistoryRecord.inspection_result = eval(iHistoryRecord.inspection_result)
            serializer = self.get_serializer(iHistoryRecord, many=False)
            if system_id is not None and system_id=='2':
                dataMap = {}
                dataMap['ret'] = 0
                dataMap['msg'] = "成功"
                dataMap['results'] = serializer.data
                dataMap['results']['video_evidence_information'] = serializer.data['inspection_result']['video_evidence_information']
                dataMap['results']['inspection_result'] = {}
                return Response(dataMap)
            else:
                dataMap = {}
                dataMap['ret'] = 0
                dataMap['msg'] = "成功"
                dataMap['results'] = serializer.data
                return Response(dataMap)

        if group_type is not None:
            if int(group_type) == 0:
                channel_ids = [1, 2, 3, 5, 6, 7, 8, 9, 10, 11, 12, 13, 15]
                conditions['channel_id__in'] = channel_ids
            elif int(group_type) == 1:
                # channel_ids = [2, 3]
                # conditions['channel_id__in'] = channel_ids
                conditions['file_type'] = 1
            elif int(group_type) == 2:
                channel_ids = [5, 6, 7, 8, 9, 10, 11, 12, 13]
                conditions['channel_id__in'] = channel_ids
            elif int(group_type) == 3:
                conditions['file_type'] = 2
            else:
                channel_ids = [4, 14, 99]
                conditions['channel_id__in'] = channel_ids

        if is_group is not None and is_group == 'true':
            historygroups = HistoryRecord.objects.filter().extra(
                {'day': 'cast(upload_time as date)'}).values_list('day').annotate(Count('id')).order_by('-day')
            results = []
            dataMap = {}
            dataMap['ret'] = 0
            dataMap['msg'] = "成功"
            dataMap['results'] = {}
            if historygroups.exists():
                if query_date is not None and int(query_date) > 0:
                    group_index = 0
                    for historygroup in historygroups:
                        if group_index < int(query_date) and group_index < len(historygroups):
                            result = {}
                            historydate = historygroup[0]
                            result["upload_time"] = historydate
                            conditions["upload_time__date"] = datetime.date(historydate.year, historydate.month, historydate.day)
                            historylist = HistoryRecord.objects.filter(
                                **conditions).order_by('-upload_time')
                            serializer_group = self.get_serializer(
                                historylist, many=True)
                            if len(serializer_group.data) > 0:
                                result["upload_datas"] = serializer_group.data
                                results.append(result)
                            group_index += 1
                    dataMap['results'] = results
                    return Response(dataMap)
                else:
                    for historygroup in historygroups:
                        result = {}
                        historydate = historygroup[0]
                        result["upload_time"] = historydate
                        conditions["upload_time__date"] = datetime.date(historydate.year, historydate.month, historydate.day)
                        historylist = HistoryRecord.objects.filter(
                            **conditions).order_by('-upload_time')
                        serializer_group = self.get_serializer(
                            historylist, many=True)
                        if len(serializer_group.data) > 0:
                            result["upload_datas"] = serializer_group.data
                            results.append(result)

                    dataMap['results'] = results
                    return Response(dataMap)
            else:
                return Response(dataMap)
        else:
            queryset = HistoryRecord.objects.filter(**conditions).order_by('-upload_time')
            #如果是pc端口访问，去掉inspection_result字段
            if system_id == '1' and objecId is None:
                for queryone in queryset:
                    queryone.inspection_result = {}

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
